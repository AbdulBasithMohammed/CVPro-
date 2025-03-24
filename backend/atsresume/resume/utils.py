import pdfplumber
import docx
import re
import spacy
import google.generativeai as genai
import phonenumbers
import json
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Gemini API with hardcoded key
genai.configure(api_key="AIzaSyBZE_RvBLkCDnADNKxstac1uv8VawmGMN8")

def extract_text_from_pdf(pdf_file):
    """
    Extract text from a PDF file using pdfplumber
    """
    try:
        text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for i, page in enumerate(pdf.pages):
                print(f"Processing page {i+1} of {len(pdf.pages)}")
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    print(f"Warning: No text extracted from page {i+1}")
        
        if not text.strip():
            print("Warning: No text extracted from the entire PDF")
            
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def parse_resume_text(text):
    """Basic parsing to extract resume fields (can be improved with NLP)"""
    resume_data = {
        "name": None,
        "email": None,
        "phone": None,
        "summary": None,
        "Experience": [],
        "skills": [],
        "projects": []
    }

    lines = text.split("\n")
    for line in lines:
        if "@" in line:  # Extract Email
            resume_data["email"] = line.strip()
        elif any(char.isdigit() for char in line) and len(line) >= 10:  # Extract Phone Number
            resume_data["phone"] = line.strip()
        elif "Experience" in line:
            resume_data["Experience"].append(line.strip())
        elif "Skills" in line:
            resume_data["skills"].append(line.strip())
        elif "Projects" in line:
            resume_data["projects"].append(line.strip())
        elif resume_data["name"] is None:  # Assume the first line is the name
            resume_data["name"] = line.strip()
    
    return resume_data



nlp = spacy.load("en_core_web_sm")  # Load NLP model for text processing

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_email(text):
    """Extract email from text using regex"""
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None

def extract_phone(text):
    """Extract phone number from text using phonenumbers library"""
    matches = re.findall(r'\+?\d[\d -]{8,15}\d', text)
    for match in matches:
        try:
            phone = phonenumbers.parse(match, None)
            if phonenumbers.is_valid_number(phone):
                return phonenumbers.format_number(phone, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
        except:
            continue
    return None

def extract_name(text):
    """Extract the most probable name using NLP"""
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None

def extract_sections(text):
    """Extract structured sections from resume"""
    sections = {
        "experience": [],
        "skills": [],
        "projects": [],
    }
    
    lines = text.split("\n")
    current_section = None

    for line in lines:
        line_lower = line.lower().strip()

        if "experience" in line_lower:
            current_section = "experience"
        elif "skills" in line_lower:
            current_section = "skills"
        elif "projects" in line_lower:
            current_section = "projects"
        elif current_section:
            sections[current_section].append(line.strip())

    return sections

def parse_resume(text):
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        **extract_sections(text),
    }

def normalize_spaces(text):
    """Normalize spaces in the text to ensure proper formatting."""
    return " ".join(text.split())

def parse_resume_with_gemini(text):
    """Uses Gemini AI to extract structured data from the resume."""
    model = genai.GenerativeModel("gemini-2.0-flash")  # Ensure correct model usage
    chat = model.start_chat()  # Start a conversation

    normalized_text = normalize_spaces(text)

    prompt = f"""
    Extract the following details from this resume:
    - Name
    - Email
    - Phone Number only in 10 digits
    - Address
    - Summary (Ensure it's extracted properly. If missing, return null.)
    - Skills (as a list)
    - Experience (Job Title, Company, Start Date, End Date, Location, Description as separate tasks)
    - Projects (Title, Description as separate tasks, Technologies used)
    - Education (Institution, Graduation Date, Course, Location)

    Structure it in **valid JSON** format:
    {json.dumps({
        "personal":{
                    "name": "string",
                    "email": "string",
                    "phone": "string",
                    "address": "string",
                    "summary": "string or null",
                    },
        "skills": ["string"],
        "experience": [
            {
                "jobTitle": "string",
                "company": "string",
                "startDate": "MM/YYYY",
                "endDate": "MM/YYYY or null (if current)",
                "location": "string",
                "tasks": ["string"]  
            }
        ],
        "projects": [
            {
                "name": "string",
                "tasks": ["string"],  
                "technologies": ["string"]
            }
        ],
        "education": [
            {
                "institution": "string",
                "graduation_date": "MM/YYYY",
                "course": "string",
                "location": "string"
            }
        ]
    }, indent=5)}

    Resume Text:
    {normalized_text}
    """

    response = chat.send_message(prompt)
    response_text = response.text.strip()

    if response_text.startswith("```json") and response_text.endswith("```"):
        response_text = response_text[7:-3].strip()  

    try:
        response_dict = json.loads(response_text)
        return response_dict
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON response: {e}")
        return None

def analyze_resume_with_gemini(resume_text):
    """
    Analyzes a resume using Google's Gemini AI and returns a score and detailed feedback.
    
    Args:
        resume_text (str): The text content of the resume to analyze
        
    Returns:
        dict: A dictionary containing the score and feedback
        {
            'score': int,
            'feedback': list[str]
        }
    """
    try:
        # Initialize Gemini model
        model = genai.GenerativeModel("gemini-2.0-flash")
        chat = model.start_chat()

        # Define ATS rules and scoring criteria
        ats_rules = """
        Evaluate this resume based on the following ATS (Applicant Tracking System) rules:

        1. Contact Information (10 points):
        - Must have full name, professional email, phone number
        - LinkedIn profile is a plus
        - Professional address format

        2. Formatting and Structure (20 points):
        - Clear section headings
        - Consistent formatting (fonts, spacing)
        - Proper use of bullet points
        - No complex tables or graphics
        - PDF format compatibility

        3. Content Organization (25 points):
        - Clear chronological or functional structure
        - Work experience with dates
        - Education with dates
        - Skills section present
        - Achievements quantified with metrics

        4. Keywords and Language (25 points):
        - Industry-specific keywords present
        - Action verbs at the start of bullets
        - Technical skills relevant to job market
        - No jargon or abbreviations
        - Proper grammar and spelling

        5. Professional Impact (20 points):
        - Achievements clearly stated
        - Impact metrics (percentages, numbers)
        - Leadership or initiative examples
        - Project outcomes
        - Awards or recognitions

        Analyze the resume and provide:
        1. A score out of 100
        2. A list of strengths (what rules were followed well)
        3. A list of improvements needed (what rules were violated)
        4. Brief explanation for each point

        Format the response as a JSON object with the following structure:
        {
            "score": number,
            "strengths": [
                {"rule": "string", "explanation": "string"}
            ],
            "improvements": [
                {"rule": "string", "explanation": "string"}
            ]
        }
        """

        # Combine resume text with ATS rules for the prompt
        prompt = f"""
        Here is the resume content to analyze:

        {resume_text}

        {ats_rules}
        """

        # Get response from Gemini
        response = chat.send_message(prompt)
        response_text = response.text.strip()

        # Extract JSON from response
        json_start = response_text.find('{')
        if json_start == -1:
            raise ValueError('Invalid response format from AI')

        json_end = response_text.rindex('}') + 1
        json_str = response_text[json_start:json_end]

        # Parse the JSON response
        result = json.loads(json_str)

        # Format the feedback
        formatted_feedback = []

        # Add strengths with checkmarks
        if result.get('strengths'):
            for strength in result['strengths']:
                if isinstance(strength, dict):
                    formatted_feedback.append(f"✓ {strength['rule']}: {strength['explanation']}")
                elif isinstance(strength, str):
                    formatted_feedback.append(f"✓ {strength}")

        # Add improvements with X marks
        if result.get('improvements'):
            for improvement in result['improvements']:
                if isinstance(improvement, dict):
                    formatted_feedback.append(f"✗ {improvement['rule']}: {improvement['explanation']}")
                elif isinstance(improvement, str):
                    formatted_feedback.append(f"✗ {improvement}")

        return {
            'score': result.get('score', 0),
            'feedback': formatted_feedback
        }

    except json.JSONDecodeError as e:
        print(f"Error parsing JSON response: {e}")
        raise ValueError("Failed to parse AI response")
    except Exception as e:
        print(f"Error analyzing resume: {e}")
        raise Exception(f"Failed to analyze resume: {str(e)}")
